from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import pandas as pd
import json
import uuid
import logging
import threading
import time
from datetime import datetime
from werkzeug.utils import secure_filename
import traceback
import sys
import re
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import PyPDF2
import docx
import spacy

# Import custom modules
import config
from models import init_db, Document, DocumentChunk, Job, JobResult, GeneratedDocument, Session
from vector_db.document_processor import DocumentProcessor
from vector_db.vector_store import VectorStore
from llm_agent.agent import LLMAgent
from llm_agent.ollama_client import OllamaClient
from document_generator.generator import DocumentGenerator

# Download NLTK resources
nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)
nltk.download('wordnet', quiet=True)

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except:
    # If model not found, download it
    os.system("python -m spacy download en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize Flask app
app = Flask(__name__)
CORS(app)

# Initialize database
init_db()

# Create upload folder
if not os.path.exists(config.UPLOAD_FOLDER):
    os.makedirs(config.UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = config.UPLOAD_FOLDER

# Create output folder
if not os.path.exists(config.OUTPUT_FOLDER):
    os.makedirs(config.OUTPUT_FOLDER)

# Skills dictionary for matching
SKILLS = {
    'technical': [
        'python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin',
        'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'spring',
        'sql', 'nosql', 'mongodb', 'mysql', 'postgresql', 'oracle', 'firebase',
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'github',
        'machine learning', 'deep learning', 'ai', 'data science', 'big data',
        'hadoop', 'spark', 'tableau', 'power bi', 'excel', 'vba',
        'html', 'css', 'sass', 'less', 'bootstrap', 'tailwind',
        'rest api', 'graphql', 'microservices', 'serverless',
        'linux', 'unix', 'windows', 'macos', 'android', 'ios',
        'agile', 'scrum', 'kanban', 'jira', 'confluence'
    ],
    'soft': [
        'communication', 'teamwork', 'leadership', 'problem solving', 'critical thinking',
        'time management', 'organization', 'adaptability', 'flexibility', 'creativity',
        'work ethic', 'interpersonal skills', 'emotional intelligence', 'conflict resolution',
        'decision making', 'stress management', 'attention to detail', 'customer service',
        'presentation', 'negotiation', 'persuasion', 'mentoring', 'coaching'
    ],
    'business': [
        'marketing', 'sales', 'finance', 'accounting', 'hr', 'human resources',
        'project management', 'product management', 'operations', 'strategy',
        'business development', 'customer relationship management', 'crm',
        'supply chain', 'logistics', 'procurement', 'quality assurance', 'qa',
        'business analysis', 'data analysis', 'market research', 'competitive analysis',
        'budgeting', 'forecasting', 'risk management', 'compliance', 'legal'
    ]
}

def allowed_file(filename, file_type):
    """Check if file has an allowed extension"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in config.ALLOWED_EXTENSIONS.get(file_type, [])

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file"""
    text = ""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
    return text

def extract_text_from_docx(docx_file):
    """Extract text from DOCX file"""
    text = ""
    try:
        doc = docx.Document(docx_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
    return text

def preprocess_text(text):
    """Preprocess text for analysis"""
    # Convert to lowercase
    text = text.lower()
    
    # Remove special characters and numbers
    text = re.sub(r'[^a-zA-Z\s]', '', text)
    
    # Tokenize
    tokens = word_tokenize(text)
    
    # Remove stopwords
    stop_words = set(stopwords.words('english'))
    tokens = [word for word in tokens if word not in stop_words]
    
    # Lemmatize
    lemmatizer = WordNetLemmatizer()
    tokens = [lemmatizer.lemmatize(word) for word in tokens]
    
    return ' '.join(tokens)

def extract_skills(text):
    """Extract skills from text"""
    skills = {
        'technical': [],
        'soft': [],
        'business': []
    }
    
    # Process with spaCy for better entity recognition
    doc = nlp(text.lower())
    
    # Extract skills from each category
    for category, skill_list in SKILLS.items():
        for skill in skill_list:
            if skill in text.lower():
                # Check if it's a complete word or part of a larger word
                pattern = r'\b' + re.escape(skill) + r'\b'
                if re.search(pattern, text.lower()):
                    skills[category].append(skill)
    
    return skills

def calculate_match_score(student_skills, job_skills):
    """Calculate match score between student skills and job skills"""
    if not job_skills:
        return 0
    
    # Flatten skills lists
    student_skills_flat = []
    for category in student_skills:
        student_skills_flat.extend(student_skills[category])
    
    job_skills_flat = []
    for category in job_skills:
        job_skills_flat.extend(job_skills[category])
    
    # Calculate match score
    if not job_skills_flat:
        return 0
    
    matches = set(student_skills_flat).intersection(set(job_skills_flat))
    score = len(matches) / len(job_skills_flat) * 100
    
    return min(score, 100)  # Cap at 100%

def analyze_student_profile(student_data):
    """Analyze student profile and extract skills"""
    # Combine all relevant fields
    profile_text = ""
    for key, value in student_data.items():
        if isinstance(value, str):
            profile_text += value + " "
    
    # Preprocess text
    processed_text = preprocess_text(profile_text)
    
    # Extract skills
    skills = extract_skills(processed_text)
    
    return skills

def analyze_job_description(job_desc_text):
    """Analyze job description and extract skills"""
    # Preprocess text
    processed_text = preprocess_text(job_desc_text)
    
    # Extract skills
    skills = extract_skills(processed_text)
    
    return skills

def process_files(job_id, csv_path, job_desc_path):
    """Process uploaded files and generate personalized documents"""
    session = Session()
    try:
        # Get job from database
        job = session.query(Job).filter(Job.job_id == job_id).first()
        if not job:
            logger.error(f"Job {job_id} not found")
            return
        
        # Update job status
        job.status = 'processing'
        session.commit()
        
        # Initialize document processor
        doc_processor = DocumentProcessor()
        
        # Process CSV file
        try:
            # Read CSV file
            df = pd.read_csv(csv_path)
            student_data = df.to_dict('records')
            job.total_students = len(student_data)
            session.commit()
            
            # Read job description file
            job_desc_text = ""
            job_desc_filename = os.path.basename(job_desc_path)
            if job_desc_filename.endswith('.pdf'):
                with open(job_desc_path, 'rb') as f:
                    job_desc_text = extract_text_from_pdf(f)
            elif job_desc_filename.endswith('.docx'):
                with open(job_desc_path, 'rb') as f:
                    job_desc_text = extract_text_from_docx(f)
            else:
                with open(job_desc_path, 'r') as f:
                    job_desc_text = f.read()
            
            # Analyze job description
            job_skills = analyze_job_description(job_desc_text)
            
            # Process each student
            for i, student in enumerate(student_data):
                # Update progress
                job.processed_students = i + 1
                session.commit()
                
                # Get student name and email
                student_name = student.get('full name of the student', student.get('Name', student.get('name', f'Student {i+1}')))
                student_email = student.get('email', student.get('Email', ''))
                student_id = student.get('Roll_Number', student.get('RollNumber', student.get('ID', student.get('StudentID', f'S{1000+i}'))))
                
                # Analyze student profile
                student_skills = analyze_student_profile(student)
                
                # Calculate match score
                match_score = calculate_match_score(student_skills, job_skills)
                
                # Determine status based on match score
                if match_score >= 70:
                    status = 'Success'
                elif match_score >= 40:
                    status = 'Partial Success'
                else:
                    status = 'Failure'
                
                # Initialize LLM agent
                llm_agent = LLMAgent()
                
                # Extract company and role from job description filename
                parts = os.path.splitext(job_desc_filename)[0].split('_')
                company = parts[0] if len(parts) > 0 else "Unknown"
                role = parts[1] if len(parts) > 1 else "Unknown"
                
                # Generate personalized document
                document_content = llm_agent.generate_personalized_document(
                    student_email=student_email,
                    company=company,
                    role=role
                )
                
                # Initialize document generator
                doc_generator = DocumentGenerator()
                
                # Generate documents in all formats
                markdown_result = doc_generator.generate_document(
                    content=document_content,
                    student_email=student_email,
                    company=company,
                    role=role,
                    job_id=job.id,
                    format_type='markdown'
                )
                
                pdf_result = doc_generator.generate_document(
                    content=document_content,
                    student_email=student_email,
                    company=company,
                    role=role,
                    job_id=job.id,
                    format_type='pdf'
                )
                
                docx_result = doc_generator.generate_document(
                    content=document_content,
                    student_email=student_email,
                    company=company,
                    role=role,
                    job_id=job.id,
                    format_type='docx'
                )
                
                # Create job result
                result = JobResult(
                    job_id=job.id,
                    student_name=student_name,
                    student_email=student_email,
                    student_id=student_id,
                    status=status,
                    match_score=match_score,
                    skills=student_skills
                )
                session.add(result)
                session.commit()
                
                # Simulate processing time
                time.sleep(0.5)
        except Exception as e:
            logger.error(f"Error processing files: {e}")
            traceback.print_exc()
            job.status = 'failed'
            session.commit()
            return
        
        # Update job status
        job.status = 'completed'
        job.completed_at = datetime.utcnow()
        session.commit()
        
    except Exception as e:
        logger.error(f"Error in process_files: {e}")
        traceback.print_exc()
        try:
            job.status = 'failed'
            session.commit()
        except:
            pass
    finally:
        session.close()

@app.route('/api/status', methods=['GET'])
def status():
    """API status endpoint"""
    try:
        # Check if Ollama is available
        ollama_client = OllamaClient()
        models = ollama_client.list_models()
        
        return jsonify({
            "status": "Server is running",
            "ollama_status": "connected" if models else "disconnected",
            "available_models": models,
            "current_model": config.OLLAMA_MODEL,
            "alternative_model": config.OLLAMA_ALTERNATIVE_MODEL
        })
    except Exception as e:
        logger.error(f"Error checking status: {e}")
        return jsonify({
            "status": "Server is running",
            "ollama_status": "error",
            "error": str(e)
        })

@app.route('/api/upload', methods=['POST'])
def upload_files():
    """Upload files endpoint"""
    try:
        # Check if files are present
        if 'csv' not in request.files or 'jobDesc' not in request.files:
            return jsonify({"error": "Missing files"}), 400
        
        csv_file = request.files['csv']
        job_desc_file = request.files['jobDesc']
        
        # Check if files are valid
        if csv_file.filename == '' or job_desc_file.filename == '':
            return jsonify({"error": "No selected files"}), 400
        
        # Check file extensions
        if not allowed_file(csv_file.filename, 'csv'):
            return jsonify({"error": "Invalid CSV file format"}), 400
        
        if not allowed_file(job_desc_file.filename, 'document'):
            return jsonify({"error": "Invalid job description file format"}), 400
        
        # Create database session
        session = Session()
        
        try:
            # Save CSV file
            csv_filename = secure_filename(csv_file.filename)
            csv_path = os.path.join(app.config['UPLOAD_FOLDER'], csv_filename)
            csv_file.save(csv_path)
            
            # Save job description file
            job_desc_filename = secure_filename(job_desc_file.filename)
            job_desc_path = os.path.join(app.config['UPLOAD_FOLDER'], job_desc_filename)
            job_desc_file.save(job_desc_path)
            
            # Create document records
            csv_document = Document(
                filename=csv_filename,
                original_filename=csv_file.filename,
                file_type='csv',
                document_type='student_data',
                file_path=csv_path,
                upload_date=datetime.utcnow()
            )
            session.add(csv_document)
            
            job_desc_document = Document(
                filename=job_desc_filename,
                original_filename=job_desc_file.filename,
                file_type=job_desc_filename.split('.')[-1].lower(),
                document_type='job_description',
                file_path=job_desc_path,
                upload_date=datetime.utcnow()
            )
            session.add(job_desc_document)
            
            # Create job ID
            job_id = str(uuid.uuid4())
            
            # Create job record
            job = Job(
                job_id=job_id,
                status='created',
                created_at=datetime.utcnow(),
                total_students=0,
                processed_students=0,
                job_description_length=os.path.getsize(job_desc_path)
            )
            session.add(job)
            session.commit()
            
            # Associate documents with job
            job.documents.append(csv_document)
            job.documents.append(job_desc_document)
            session.commit()
            
            # Process documents in the background
            doc_processor = DocumentProcessor()
            doc_processor.process_document(csv_document.id)
            doc_processor.process_document(job_desc_document.id)
            
            # Start processing in background
            threading.Thread(
                target=process_files,
                args=(job_id, csv_path, job_desc_path)
            ).start()
            
            return jsonify({
                "job_id": job_id,
                "status": "processing",
                "message": "Files uploaded successfully and processing started"
            })
            
        except Exception as e:
            session.rollback()
            logger.error(f"Error in upload_files: {e}")
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Error uploading files: {e}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/api/job/<job_id>', methods=['GET'])
def get_job_status(job_id):
    """Get job status endpoint"""
    session = Session()
    try:
        # Get job from database
        job = session.query(Job).filter(Job.job_id == job_id).first()
        if not job:
            return jsonify({"error": "Job not found"}), 404
        
        # Get job results
        results = []
        if job.status == 'completed':
            for result in job.results:
                results.append({
                    'name': result.student_name,
                    'email': result.student_email,
                    'rollNumber': result.student_id,
                    'status': result.status,
                    'matchScore': result.match_score,
                    'skills': result.skills
                })
        
        # Get generated documents
        documents = []
        for doc in job.generated_documents:
            documents.append({
                'id': doc.id,
                'student_email': doc.student_email,
                'company': doc.company,
                'role': doc.role,
                'document_type': doc.document_type,
                'file_path': doc.file_path,
                'generated_at': doc.generated_at.isoformat()
            })
        
        response = {
            "job_id": job_id,
            "status": job.status,
            "created_at": job.created_at.isoformat(),
            "completed_at": job.completed_at.isoformat() if job.completed_at else None,
            "total_students": job.total_students,
            "processed_students": job.processed_students,
            "results": results,
            "documents": documents
        }
        
        return jsonify(response)
    except Exception as e:
        logger.error(f"Error getting job status: {e}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/document/<int:document_id>', methods=['GET'])
def get_document(document_id):
    """Get generated document endpoint"""
    session = Session()
    try:
        # Get document from database
        document = session.query(GeneratedDocument).filter(GeneratedDocument.id == document_id).first()
        if not document:
            return jsonify({"error": "Document not found"}), 404
        
        # Check if file exists
        if not os.path.exists(document.file_path):
            return jsonify({"error": "Document file not found"}), 404
        
        # Return file
        return send_file(
            document.file_path,
            as_attachment=True,
            download_name=os.path.basename(document.file_path)
        )
    except Exception as e:
        logger.error(f"Error getting document: {e}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/documents/formats', methods=['GET'])
def get_document_formats():
    """Get available document formats endpoint"""
    return jsonify({
        "formats": [
            {"id": "markdown", "name": "Markdown (.md)"},
            {"id": "pdf", "name": "PDF Document (.pdf)"},
            {"id": "docx", "name": "Word Document (.docx)"}
        ]
    })

@app.route('/api/student/<email>/documents', methods=['GET'])
def get_student_documents(email):
    """Get documents for a student endpoint"""
    session = Session()
    try:
        # Get documents from database
        documents = session.query(GeneratedDocument).filter(GeneratedDocument.student_email == email).all()
        
        # Format response
        results = []
        for doc in documents:
            results.append({
                'id': doc.id,
                'student_email': doc.student_email,
                'company': doc.company,
                'role': doc.role,
                'document_type': doc.document_type,
                'file_path': doc.file_path,
                'generated_at': doc.generated_at.isoformat()
            })
        
        return jsonify({"documents": results})
    except Exception as e:
        logger.error(f"Error getting student documents: {e}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/generate/document', methods=['POST'])
def generate_document():
    """Generate document endpoint"""
    try:
        # Get request data
        data = request.json
        if not data:
            return jsonify({"error": "No data provided"}), 400
        
        # Check required fields
        required_fields = ['student_email', 'company', 'role', 'job_id', 'format']
        for field in required_fields:
            if field not in data:
                return jsonify({"error": f"Missing required field: {field}"}), 400
        
        # Get job
        session = Session()
        try:
            job = session.query(Job).filter(Job.job_id == data['job_id']).first()
            if not job:
                return jsonify({"error": "Job not found"}), 404
            
            # Initialize LLM agent
            llm_agent = LLMAgent()
            
            # Generate personalized document
            document_content = llm_agent.generate_personalized_document(
                student_email=data['student_email'],
                company=data['company'],
                role=data['role']
            )
            
            # Initialize document generator
            doc_generator = DocumentGenerator()
            
            # Generate document in requested format
            result = doc_generator.generate_document(
                content=document_content,
                student_email=data['student_email'],
                company=data['company'],
                role=data['role'],
                job_id=job.id,
                format_type=data['format']
            )
            
            if 'error' in result:
                return jsonify({"error": result['error']}), 500
            
            return jsonify({
                "document_id": result['document_id'],
                "file_path": result['file_path'],
                "document_type": result['document_type'],
                "message": f"Document generated successfully in {result['document_type']} format"
            })
        except Exception as e:
            logger.error(f"Error generating document: {e}")
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500
        finally:
            session.close()
    except Exception as e:
        logger.error(f"Error in generate_document: {e}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=config.PORT, debug=config.DEBUG)