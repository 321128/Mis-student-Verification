import os
import sys
import logging
import markdown2
import pdfkit
from docx import Document as DocxDocument
from datetime import datetime
from typing import Dict, Any, List, Optional

# Add parent directory to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import config
from models import GeneratedDocument, Session

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Generate documents in different formats from markdown content"""
    
    def __init__(self):
        """Initialize the document generator"""
        self.output_folder = config.OUTPUT_FOLDER
        
        # Create output folder if it doesn't exist
        if not os.path.exists(self.output_folder):
            os.makedirs(self.output_folder)
    
    def generate_document(
        self, 
        content: str, 
        student_email: str, 
        company: str, 
        role: str, 
        job_id: int,
        format_type: str = 'markdown'
    ) -> Dict[str, Any]:
        """Generate a document in the specified format
        
        Args:
            content: Markdown content
            student_email: Email of the student
            company: Company name
            role: Job role
            job_id: ID of the job
            format_type: Format type ('markdown', 'pdf', 'docx')
            
        Returns:
            Dictionary with file path and metadata
        """
        # Sanitize filename components
        student_email = self._sanitize_filename(student_email)
        company = self._sanitize_filename(company)
        role = self._sanitize_filename(role)
        
        # Create base filename
        base_filename = f"{student_email}_{company}_{role}"
        
        # Generate document based on format type
        if format_type == 'markdown':
            return self._generate_markdown(content, base_filename, job_id)
        elif format_type == 'pdf':
            return self._generate_pdf(content, base_filename, job_id)
        elif format_type == 'docx':
            return self._generate_docx(content, base_filename, job_id)
        else:
            logger.error(f"Unsupported format type: {format_type}")
            return {"error": f"Unsupported format type: {format_type}"}
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename to remove invalid characters
        
        Args:
            filename: Filename to sanitize
            
        Returns:
            Sanitized filename
        """
        # Replace invalid characters with underscores
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        
        # Limit length
        return filename[:50]
    
    def _generate_markdown(self, content: str, base_filename: str, job_id: int) -> Dict[str, Any]:
        """Generate a markdown document
        
        Args:
            content: Markdown content
            base_filename: Base filename
            job_id: ID of the job
            
        Returns:
            Dictionary with file path and metadata
        """
        try:
            # Create filename
            filename = f"{base_filename}.md"
            file_path = os.path.join(self.output_folder, filename)
            
            # Write content to file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            # Save to database
            session = Session()
            try:
                # Extract email from filename
                email = base_filename.split('_')[0]
                
                # Extract company and role
                parts = base_filename.split('_')
                company = parts[1] if len(parts) > 1 else ""
                role = parts[2] if len(parts) > 2 else ""
                
                # Create document record
                doc = GeneratedDocument(
                    job_id=job_id,
                    student_email=email,
                    company=company,
                    role=role,
                    document_type='markdown',
                    file_path=file_path,
                    content=content
                )
                session.add(doc)
                session.commit()
                
                logger.info(f"Generated markdown document: {file_path}")
                return {
                    "file_path": file_path,
                    "document_type": 'markdown',
                    "document_id": doc.id
                }
            except Exception as e:
                session.rollback()
                logger.error(f"Error saving document to database: {e}")
                return {
                    "file_path": file_path,
                    "document_type": 'markdown',
                    "error": str(e)
                }
            finally:
                session.close()
        except Exception as e:
            logger.error(f"Error generating markdown document: {e}")
            return {"error": str(e)}
    
    def _generate_pdf(self, content: str, base_filename: str, job_id: int) -> Dict[str, Any]:
        """Generate a PDF document
        
        Args:
            content: Markdown content
            base_filename: Base filename
            job_id: ID of the job
            
        Returns:
            Dictionary with file path and metadata
        """
        try:
            # Convert markdown to HTML
            html = markdown2.markdown(content)
            
            # Create filename
            filename = f"{base_filename}.pdf"
            file_path = os.path.join(self.output_folder, filename)
            
            # Convert HTML to PDF
            pdfkit.from_string(html, file_path)
            
            # Save to database
            session = Session()
            try:
                # Extract email from filename
                email = base_filename.split('_')[0]
                
                # Extract company and role
                parts = base_filename.split('_')
                company = parts[1] if len(parts) > 1 else ""
                role = parts[2] if len(parts) > 2 else ""
                
                # Create document record
                doc = GeneratedDocument(
                    job_id=job_id,
                    student_email=email,
                    company=company,
                    role=role,
                    document_type='pdf',
                    file_path=file_path,
                    content=content  # Store original markdown
                )
                session.add(doc)
                session.commit()
                
                logger.info(f"Generated PDF document: {file_path}")
                return {
                    "file_path": file_path,
                    "document_type": 'pdf',
                    "document_id": doc.id
                }
            except Exception as e:
                session.rollback()
                logger.error(f"Error saving document to database: {e}")
                return {
                    "file_path": file_path,
                    "document_type": 'pdf',
                    "error": str(e)
                }
            finally:
                session.close()
        except Exception as e:
            logger.error(f"Error generating PDF document: {e}")
            return {"error": str(e)}
    
    def _generate_docx(self, content: str, base_filename: str, job_id: int) -> Dict[str, Any]:
        """Generate a DOCX document
        
        Args:
            content: Markdown content
            base_filename: Base filename
            job_id: ID of the job
            
        Returns:
            Dictionary with file path and metadata
        """
        try:
            # Create filename
            filename = f"{base_filename}.docx"
            file_path = os.path.join(self.output_folder, filename)
            
            # Create document
            doc = DocxDocument()
            
            # Parse markdown and add to document
            # This is a simple implementation - for production, use a more robust markdown to docx converter
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Handle headings
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                # Handle bullet points
                elif line.startswith('* ') or line.startswith('- '):
                    doc.add_paragraph(line[2:], style='ListBullet')
                # Handle numbered lists
                elif line.startswith('1. ') or line.startswith('1) '):
                    doc.add_paragraph(line[3:], style='ListNumber')
                # Handle regular paragraphs
                else:
                    doc.add_paragraph(line)
            
            # Save document
            doc.save(file_path)
            
            # Save to database
            session = Session()
            try:
                # Extract email from filename
                email = base_filename.split('_')[0]
                
                # Extract company and role
                parts = base_filename.split('_')
                company = parts[1] if len(parts) > 1 else ""
                role = parts[2] if len(parts) > 2 else ""
                
                # Create document record
                doc = GeneratedDocument(
                    job_id=job_id,
                    student_email=email,
                    company=company,
                    role=role,
                    document_type='docx',
                    file_path=file_path,
                    content=content  # Store original markdown
                )
                session.add(doc)
                session.commit()
                
                logger.info(f"Generated DOCX document: {file_path}")
                return {
                    "file_path": file_path,
                    "document_type": 'docx',
                    "document_id": doc.id
                }
            except Exception as e:
                session.rollback()
                logger.error(f"Error saving document to database: {e}")
                return {
                    "file_path": file_path,
                    "document_type": 'docx',
                    "error": str(e)
                }
            finally:
                session.close()
        except Exception as e:
            logger.error(f"Error generating DOCX document: {e}")
            return {"error": str(e)}